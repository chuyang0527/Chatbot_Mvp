"""文档处理模块"""

from pathlib import Path
from typing import List
import pypdf
from docx import Document
import markdown

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangChainDocument

from config import Config


class DocumentProcessor:
    """文档处理器，支持多种格式"""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=Config.CHUNK_SIZE,
            chunk_overlap=Config.CHUNK_OVERLAP,
            length_function=len,
        )
    
    def load_document(self, file_path: str) -> List[LangChainDocument]:
        """加载并处理文档"""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 根据文件类型提取文本
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            text = self._extract_pdf(path)
        elif suffix == ".docx":
            text = self._extract_docx(path)
        elif suffix == ".md":
            text = self._extract_markdown(path)
        elif suffix == ".txt":
            text = self._extract_txt(path)
        else:
            raise ValueError(f"不支持的文件格式: {suffix}")
        
        # 分段
        chunks = self.text_splitter.split_text(text)
        
        # 创建文档对象
        documents = [
            LangChainDocument(
                page_content=chunk,
                metadata={
                    "source": str(path.name),
                    "chunk_id": i,
                    "total_chunks": len(chunks)
                }
            )
            for i, chunk in enumerate(chunks)
        ]
        
        return documents
    
    def _extract_pdf(self, path: Path) -> str:
        """提取 PDF 文本"""
        text = []
        with open(path, "rb") as file:
            pdf_reader = pypdf.PdfReader(file)
            for page in pdf_reader.pages:
                text.append(page.extract_text())
        return "\n\n".join(text)
    
    def _extract_docx(self, path: Path) -> str:
        """提取 Word 文本"""
        doc = Document(path)
        return "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    
    def _extract_markdown(self, path: Path) -> str:
        """提取 Markdown 文本"""
        with open(path, "r", encoding="utf-8") as file:
            md_text = file.read()
        # 转换为纯文本（去除 Markdown 标记）
        html = markdown.markdown(md_text)
        # 简单去除 HTML 标签
        import re
        text = re.sub(r"<[^>]+>", "", html)
        return text
    
    def _extract_txt(self, path: Path) -> str:
        """提取 TXT 文本"""
        with open(path, "r", encoding="utf-8") as file:
            return file.read()
